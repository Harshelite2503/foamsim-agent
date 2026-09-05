"""Build FoamSim_Agent_Research_Proposal.docx from live results."""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
ROOT = Path(__file__).resolve().parent.parent
g = pd.read_csv(ROOT / "runs/grades.csv"); g = g[g.has_script == 1]
n = len(g); n_phys = int(g.physics_pass.sum()); n_pb = int(g[g.workflow == "W4_ill_posed"].rubric_pushback.sum())
LEVELS = ["L1_sketch", "L2_goal", "L3_recipe", "L4_spec", "L5_contract"]
it = g.groupby("level").meta_iterations.mean().reindex(LEVELS); loc = g.groupby("level").loc.mean().reindex(LEVELS); ru = g.groupby("level").reusable.mean().reindex(LEVELS)
val = pd.read_csv(ROOT / "data/validation.csv"); k = val[val.grade == "K46"].dropna(subset=["E_FE_mean"])
doc = Document(); doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(11)
for s in doc.sections: s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1)
def H(t, l=1): doc.add_heading(t, level=l)
def P(t, italic=False, align=None):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic
    if align is not None: p.alignment = align
def B(items, style="List Bullet"):
    for i in items: doc.add_paragraph(i, style=style)
def T(header, body, widths=None):
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]; c.text = ""; c.paragraphs[0].add_run(h).bold = True
    for row in body:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = str(v)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Inches(w)
    doc.add_paragraph()
def FIG(path, cap, w=6.0):
    doc.add_picture(str(path), width=Inches(w)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER; P(cap, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("FoamSim-Agent"); r.bold = True; r.font.size = Pt(24)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run("AI Coding Agents for Trustworthy Micromechanics Simulation of Hollow-Particle Composites: Toolkit, Agent Skills, and a Physics-Graded Benchmark"); r.font.size = Pt(14); r.italic = True
P("Research Proposal", align=WD_ALIGN_PARAGRAPH.CENTER)
P("Harsh Vardhan Gupta (proposer)  ·  Prof. Nikhil Gupta, FASM (faculty collaborator)\nDepartment of Mechanical and Aerospace Engineering, NYU Tandon School of Engineering", align=WD_ALIGN_PARAGRAPH.CENTER)
P("September 2026  ·  Code: github.com/Harshelite2503/foamsim-agent", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

H("Abstract")
P(f"AI coding agents now turn a sentence into a runnable simulation. NVIDIA's ALCHEMI study showed this for atomistic simulation and reported two cautions: prompt detail changes code structure but not physics, and agents never question an ill-posed request. Composite micromechanics — the field of the collaborating group — has neither an agent-ready toolkit nor a benchmark, and its ground truth is experimental rather than ab initio. We propose FoamSim-Agent: a small validated simulation toolkit for syntactic foams with built-in physical admissibility guards, agent skills documenting its API and validation protocol, and a benchmark of four workflows across five prompt-specificity levels, graded on code features, execution and physics agreement with analytical references and the FoamGPT experimental dataset. A pilot of {n} agent-generated pipelines is complete: {n_phys}/{n} executed and passed the physics checks; {n_pb}/5 ill-posed requests were refused, each by way of the toolkit's guards — inverting ALCHEMI's zero-pushback result and isolating library-level checks as the enforceable lever; specificity raised reusability from {ru['L1_sketch']:.1f} to {ru['L5_contract']:.1f} at {it['L5_contract']/it['L1_sketch']:.1f}× the iterations. All five finite-element runs correctly diagnosed a 0.9% overshoot of the Hashin–Shtrikman bound as finite-cell KUBC bias rather than tuning it away. The proposal scales the pilot to multiple agent models and samples, adds periodic-boundary FE and a second ill-posed control, and anchors the benchmark with like-for-like laboratory measurements.")

H("1. Problem Statement")
H("1.1 Simulation is gated by software skill, not by science", 2)
P("Micromechanics of syntactic foams is mature — Hashin–Shtrikman bounds, Mori–Tanaka and differential schemes, Bardella–Genna and Porfiri–Gupta models, finite-element unit cells — yet most experimentalists in the field do not run these models because doing so requires writing and validating code. Coding agents remove that barrier, but only if the generated pipelines are physically right.")
H("1.2 Nobody has measured whether agent-generated composite simulations are trustworthy", 2)
P("Existing evaluations target atomistic simulation (ALCHEMI), generic FEM code generation (FEM-Bench), or structural design; none grades against experimental data, and none includes an ill-posed control. ALCHEMI's own finding — zero pushback on unphysical tasks across every configuration — is unresolved: is it a property of agents, or of unguarded toolkits?")
H("1.3 Experimental ground truth is harder than DFT ground truth", 2)
P("Literature moduli for polymer syntactic foams are heterogeneous in measurement method; whether they can validate micromechanics at all is unknown. A benchmark in this field must confront that.")

H("2. Objectives")
B(["O1 — Toolkit: a validated, composable, guarded micromechanics library for hollow-particle composites (analytical + FE) with an experimental data bridge.",
   "O2 — Skills: agent-loadable API patterns and a validation protocol, in the open Agent Skills format.",
   "O3 — Benchmark: workflows × prompt levels × agent models × samples, graded on code features, execution and physics, with ill-posed controls the toolkit can and cannot catch.",
   "O4 — Findings: what determines physical correctness and pushback; cost–structure trade-offs; what experimental data can validate.",
   "O5 — Anchor: laboratory measurements (epoxy/K46, strain-gauge moduli) that give the benchmark a like-for-like experimental reference."], "List Number")

H("3. Approach")
T(["Component", "Design", "Status"], [
 ["foamsim toolkit", "HS/Walpole estimate → bounds, Mori–Tanaka, exact hollow-sphere K (Hashin CSA); HP-MT and HP-DS; Gibson–Ashby; crush onset; RSA packing; scikit-fem voxel FE (KUBC); FoamGPT bridge; guards on vf > 0.64 and η ∉ [0,1)", "Done, 13 tests"],
 ["Validation", "FE vs HP-MT/HP-DS/HS for three 3M grades; homogeneous-box exactness", "Done (Figure 1)"],
 ["Skills", "foamsim-micromechanics, foamsim-fe-rve, foamsim-validate", "Done"],
 ["Benchmark v0", "4 workflows × 5 levels × 1 sample, Claude Opus 5; grader (code features, execution, physics + rubric)", "Done, 20 runs"],
 ["Benchmark v1", "+ second model family, 3 samples/cell, periodic BCs, 'shell' FE runs, second ill-posed control", "Proposed"],
 ["Lab anchor", "cast epoxy/K46 at vf 0.2–0.5, strain-gauge compressive moduli", "Proposed (Gupta lab)"],
], widths=[1.3, 3.9, 1.3])

H("4. Preliminary Results")
FIG(ROOT / "data/validation.png", "Figure 1. Toolkit validation: analytical estimates, HS bands, FE-KUBC (K46) and FoamGPT experimental data.")
P("Toolkit: FE-KUBC exceeds HP-MT by " + ", ".join(f"{(r.E_FE_mean/r.E_MT-1)*100:.1f}% (vf {r.vf})" for _, r in k.iterrows()) + "; seed spread < 0.2%; homogeneous box exact to 1e-16.")
T(["Metric", "Result"], [
 ["Pipelines executed / physics-correct", f"{int(g.executed.sum())}/{n} · {n_phys}/{n}"],
 ["Ill-posed requests refused", f"{n_pb}/5 (all via toolkit ValueError)"],
 ["Mean iterations L1 → L5", f"{it['L1_sketch']:.1f} → {it['L5_contract']:.1f}"],
 ["Mean lines of code L1 → L5", f"{loc['L1_sketch']:.0f} → {loc['L5_contract']:.0f}"],
 ["Reusable interface L1 → L5", f"{ru['L1_sketch']:.1f} → {ru['L5_contract']:.1f}"],
 ["W2 inverse design", "same optimum at all 5 levels (η ≈ 0.958, vf 0.60, 0.657 g/cm³, 3515–3540 MPa)"],
 ["W3 FE vs MT", "+2.8%; 0.9% above HS upper bound diagnosed as KUBC finite-cell bias by 5/5 runs; converges with cell size"],
 ["Data finding", "FoamGPT epoxy/GMB moduli lie below the HS lower bound → validation needs strain-metadata"],
], widths=[2.6, 3.9])
P("Two benchmark-design errors (a non-binding design constraint; a protocol line that leaked the ill-posed verdict) were caught by the agents and by review and corrected before the affected cells ran.")

H("5. Work Plan")
T(["Phase", "Weeks", "Activities", "Deliverable"], [
 ["1 (done)", "—", "Toolkit, skills, validation, 20-cell pilot, grader, draft paper", "v0 repo + draft"],
 ["2", "1–3", "Second agent model; 3 samples per cell; automated rubric; cost accounting", "Benchmark v1 results"],
 ["3", "2–5", "Periodic BCs; shell-resolved FE at n ≥ 64; second ill-posed control; strength model", "Toolkit v1"],
 ["4", "3–7", "Lab anchor: cast + test epoxy/K46 series; compare to agent pipelines", "Experimental reference"],
 ["5", "6–9", "Write-up; release skills to the Agent Skills ecosystem", "Submission"],
], widths=[0.9, 0.7, 3.6, 1.2])

H("6. Expected Outputs and Venues")
T(["Output", "Venue options", "Contribution"], [
 ["Benchmark paper", "NeurIPS AI4Mat / ICLR AI4Science workshops; npj Computational Materials; Computational Materials Science", "First physics-graded, experiment-anchored benchmark of coding agents for composite simulation; the guarded-toolkit pushback result"],
 ["Toolkit + skills", "JOSS; GitHub", "Reusable agent-ready micromechanics library"],
 ["Materials note", "Composites Part B / Materials & Design", "Experimental compressive moduli vs HS bounds: a measurement-method audit"],
], widths=[1.4, 2.4, 2.7])
H("7. Resources and Budget")
T(["Item", "Estimate"], [["Agent runs (2 models × 4 workflows × 5 levels × 3 samples)", "≈ US$150–300 API, or Claude Code quota"], ["FE compute", "Laptop CPU; optional workstation for n ≥ 64"], ["Lab anchor", "Gupta lab: resin, K46, ~20 specimens, strain-gauge compression"], ["Expert time", "≈ 6–10 h review of physics rubric and lab plan"]], widths=[4.5, 2.0])
H("8. Roles")
B(["Harsh Vardhan Gupta: toolkit, skills, benchmark harness, agent runs, grading, analysis, drafting.", "Prof. Nikhil Gupta: micromechanics review (model choices, HS/KUBC interpretation), laboratory anchor experiments, domain framing, co-authorship."])
H("9. Risks")
T(["Risk", "Mitigation"], [["Single model family / small sample in pilot", "Phase 2 adds a second model and 3 samples/cell"], ["KUBC bias confounds FE-vs-bounds checks", "Periodic BCs and larger cells in Phase 3; protocol wording updated"], ["Experimental data cannot validate", "Lab anchor series with strain gauges; measurement-method audit as a separate contribution"], ["Pushback driven only by library guards", "Second ill-posed control the toolkit cannot catch isolates agent judgment"]], widths=[2.6, 3.9])
out = ROOT / "FoamSim_Agent_Research_Proposal.docx"; doc.save(out); (ROOT / "docs" / out.name).write_bytes(out.read_bytes()); print("saved", out)
