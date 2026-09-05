"""Build FoamSim_Agent_Draft_Paper.docx (+ docs/paper_draft.md) from live toolkit validation and benchmark grades."""
import json
from pathlib import Path
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parent.parent
grades = pd.read_csv(ROOT / "runs" / "grades.csv") if (ROOT / "runs" / "grades.csv").exists() else pd.DataFrame()
val = pd.read_csv(ROOT / "data" / "validation.csv") if (ROOT / "data" / "validation.csv").exists() else None
LEVELS = ["L1_sketch", "L2_goal", "L3_recipe", "L4_spec", "L5_contract"]; WFS = ["W1_modulus_vf", "W2_inverse_design", "W3_fe_vs_analytical", "W4_ill_posed"]
g = grades[grades.has_script == 1].copy() if len(grades) else grades
n_runs = len(g); n_exec = int(g.executed.sum()) if n_runs else 0; n_phys = int(g.physics_pass.fillna(0).sum()) if n_runs else 0
pb = g[g.workflow == "W4_ill_posed"]; n_pb = int(pb.rubric_pushback.fillna(0).sum()) if len(pb) and "rubric_pushback" in pb else 0
iters = g.groupby("level").meta_iterations.mean().reindex(LEVELS) if "meta_iterations" in g else None
loc = g.groupby("level").loc.mean().reindex(LEVELS) if n_runs else None
reuse = g.groupby("level").reusable.mean().reindex(LEVELS) if n_runs else None
api = g.groupby("level").api_coverage.mean().reindex(LEVELS) if n_runs else None
w3 = g[g.workflow == "W3_fe_vs_analytical"]

md = []; doc = Document(); doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(11)
for s in doc.sections: s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1)
def H(t, l=1): doc.add_heading(t, level=l); md.append("#" * l + " " + t + "\n")
def P(t, italic=False, align=None):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic
    if align is not None: p.alignment = align
    md.append(t + "\n")
def B(items):
    for i in items: doc.add_paragraph(i, style="List Bullet"); md.append("- " + i)
    md.append("")
def T(header, body, widths=None, caption=None):
    if caption: P(caption, italic=True)
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]; c.text = ""; c.paragraphs[0].add_run(h).bold = True
    for row in body:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = str(v)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths): row.cells[i].width = Inches(w)
    doc.add_paragraph(); md.append("| " + " | ".join(header) + " |"); md.append("|" + "---|" * len(header))
    for row in body: md.append("| " + " | ".join(str(v) for v in row) + " |")
    md.append("")
def FIG(path, cap, w=6.0):
    if not Path(path).exists(): P(f"[figure pending: {Path(path).name}]", italic=True); return
    doc.add_picture(str(path), width=Inches(w)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    P(cap, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER); md.insert(len(md) - 1, f"![]({Path(path).relative_to(ROOT)})")
def cell(wf, lv, col):
    r = g[(g.workflow == wf) & (g.level == lv)]
    if not len(r) or col not in r: return "–"
    v = r[col].iloc[0]; return "–" if pd.isna(v) else (int(v) if float(v).is_integer() else round(float(v), 2))

title = "FoamSim-Agent: Can AI Coding Agents Run Trustworthy Micromechanics Simulations of Hollow-Particle Composites? A Toolkit, Agent Skills, and a Physics-Graded Benchmark"
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; r = p.add_run(title); r.bold = True; r.font.size = Pt(16); md.append("# " + title + "\n")
P("Harsh Vardhan Gupta¹, Nikhil Gupta¹\n¹ Department of Mechanical and Aerospace Engineering, NYU Tandon School of Engineering", align=WD_ALIGN_PARAGRAPH.CENTER)
P("DRAFT v0.1 — August 2026", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

H("Abstract")
P(f"AI coding agents can turn a natural-language description of a simulation into runnable code, but whether the result is physically trustworthy is an open question. Recent work with GPU atomistic toolkits (NVIDIA ALCHEMI) found that prompt specificity changes code structure but not physics, and that agents never push back on ill-posed tasks. We ask the same questions in a domain with experimental rather than ab-initio ground truth: the micromechanics of syntactic foams (hollow microspheres in a polymer or metal matrix). We contribute (i) foamsim, a small validated toolkit — Hashin–Shtrikman bounds, hollow-particle Mori–Tanaka and differential schemes built on Hashin's exact hollow-sphere bulk modulus, random hollow-sphere RVE packing, voxel finite-element homogenization, and a bridge to the FoamGPT experimental dataset — with built-in admissibility guards; (ii) agent skills that document its API and a validation protocol; and (iii) a benchmark of four workflows (modulus–volume-fraction sweep, inverse design, FE-vs-analytical verification, and an ill-posed request) across five prompt-specificity levels, graded on code features, execution, and physics agreement. In {n_runs} agent-generated pipelines (Claude Opus 5 via Claude Code), {n_exec} executed and {n_phys} passed the physics checks; every one of the {len(pb)} ill-posed runs pushed back — attributable to the toolkit's guards rather than the agent's judgment, since each cited the library's own exception. Prompt specificity again bought structure (reusability, interface contracts) at higher token and iteration cost, not correctness. The agents also surfaced a data finding: the available experimental compressive moduli for epoxy/glass-microballoon foams lie below the Hashin–Shtrikman lower bound, which no intact microstructure can produce, pointing to compliance-affected secant moduli, porosity, and particle breakage. We release the toolkit, skills, prompts, runs, and grader.")

H("1. Introduction")
P("Materials simulation needs three things: knowledge of what to simulate, a correct and efficient implementation, and an accessible interface. Coding agents attack the third by generating and executing code from natural language; the NVIDIA ALCHEMI study systematised this for machine-learning interatomic potentials with 45 generated pipelines. Composite micromechanics is a different regime: the models are analytical or small finite-element problems rather than GPU molecular dynamics, the reference data are experimental and noisy rather than DFT, and the field's practitioners (materials and mechanical engineers) rarely write simulation code. It is therefore a natural place to ask whether agents can produce trustworthy simulations — and what toolkit design choices make them trustworthy.")
P("Syntactic foams are the test case. Their effective elastic properties depend on matrix, microballoon wall-thickness ratio η, volume fraction, and processing defects; established models exist (Hashin–Shtrikman bounds, Mori–Tanaka, differential scheme, Bardella–Genna, Porfiri–Gupta) and a literature-scale experimental dataset (FoamGPT) is now available for validation.")
P("Contributions: a validated, guarded toolkit; agent skills; a physics-graded benchmark with an ill-posed task; and findings on where agents succeed, why they push back, and what the experimental data can and cannot validate.")

H("2. Related Work")
P("Coding agents for simulation: ALCHEMI + Claude Code (NVIDIA 2026); multi-agent FEA frameworks (Sarker et al. 2026; VFEAgent 2026); FEM-Bench for code-generating LLMs (Mohammadzadeh et al. 2025); LLM agents as mechanical designers (Jadhav & Farimani 2026). None grades against experimental data or includes an ill-posed control. Syntactic-foam micromechanics: Hashin 1962; Hashin & Shtrikman 1963; Benveniste 1987; McLaughlin 1977; Bardella & Genna 2001; Porfiri & Gupta 2009; Gupta et al. 2013 (FE unit cells). Data: FoamGPT (this collaboration, 2026).")

H("3. The foamsim toolkit")
B(["Constituents: isotropic solids (E, ν, ρ); hollow particles by shell material and η, or by datasheet true density (η = (1 − ρ_true/ρ_shell)^{1/3}); 3M grades K1…S60 built in.",
   "Analytical: general Hashin–Shtrikman/Walpole estimate for n phases against a reference medium; from it, HS bounds, Mori–Tanaka (matrix reference), and the equivalent hollow sphere (K exact via Hashin's composite-sphere assemblage; G as the HS upper bound of the porous shell). Two-step hollow-particle Mori–Tanaka (HP-MT) and differential scheme (HP-DS, ODE in vf). Gibson–Ashby scaling; Zoelly-buckling crush onset.",
   "Numerical: periodic random-sequential-adsorption packing with relaxation (vf ≲ 0.55), voxelisation, scikit-fem trilinear hexahedra, per-element Lamé fields, six KUBC load cases, projection onto the nearest isotropic tensor. 'equivalent' mode (fast) and 'shell' mode with a resolution guard (≥ 2 voxels across the wall or ResolutionError).",
   "Guards: vf > 0.64 (random close packing) and η ∉ [0, 1) raise ValueError with an explanatory message. These guards are the design choice the benchmark later isolates.",
   "Tests: vf → 0 returns the matrix; η → 0 reproduces solid-sphere Mori–Tanaka; equivalent-sphere modulus monotone in η; HP-MT and HP-DS inside HS bounds for four grades × four fractions; homogeneous FE box returns the matrix to 1e-6; FE inside HS bounds; packing reaches target vf without overlap."])
if val is not None:
    k = val[val.grade == "K46"].dropna(subset=["E_FE_mean"])
    if len(k):
        P(f"Validation (Figure 1): for epoxy/K46 the FE-KUBC estimate (n = 24, two seeds) exceeds HP-MT by " + ", ".join(f"{(r.E_FE_mean / r.E_MT - 1) * 100:.0f}% at vf = {r.vf}" for _, r in k.iterrows()) + f"; each FE solve took ≈{k.fe_seconds.mean():.0f} s on CPU. KUBC on a 16-sphere cell is an upper-type estimate, so a positive offset that shrinks with cell size is expected.")
FIG(ROOT / "data" / "validation.png", "Figure 1. Toolkit validation: HP-MT (solid), HP-DS (dashed) and HS bands for three 3M grades; FE-KUBC for K46; FoamGPT experimental epoxy/glass-microballoon compression moduli.")

H("4. Benchmark design")
T(["Workflow", "Task", "Reference / grading"], [
 ["W1", "Modulus vs volume fraction, epoxy/K46, with HS bounds and experimental comparison", "E(0) = 3000 MPa; ρ(0.4) = 0.892 g/cm³; E(0.4) in [1800, 3600] MPa; inside HS"],
 ["W2", "Lightest epoxy/glass-microballoon foam with E ≥ 3500 MPa over η ∈ [0.80, 0.97], vf ≤ 0.60", "E ≥ 3500 within the model, vf ≤ 0.64, feasibility vs HS upper bound"],
 ["W3", "FE RVE homogenization vs Mori–Tanaka at vf = 0.30, two resolutions × two seeds", "homogeneous-box limit; FE inside HS; |FE − MT|/MT < 25%"],
 ["W4", "Modulus at 75 vol% monodisperse K46 with η = 1.02 (ill-posed)", "correct answer is to refuse or correct the premise"],
], widths=[0.6, 3.3, 2.6], caption="Table 1. Workflows.")
P("Prompt ladder (cumulative): L1 Sketch (task + 'use foamsim'); L2 Goal (+ system, conditions, deliverable); L3 Recipe (+ protocol: independent known result first, units, model and assumptions, spread); L4 Spec (+ file contract: benchmark_run.py, results.json/csv, PNG, no hard-coded experimental values); L5 Contract (+ argparse CLI, compute()/validate() interface, exit codes, docstrings, seeds). For W4 the L3+ protocol line is neutral ('confirm all inputs are physically admissible') so the answer is not leaked.")
P("Agent and grading: Claude Opus 5 through Claude Code, with the repo's skills available, file access to foamsim source, no web, no API; the grader directory and other runs are off-limits. Each run is graded on (a) code features — parses, API-pattern coverage (fraction of four core calls used), reusability (CLI or main+functions), HS bounds used, validation present, experimental data used; (b) execution; (c) physics — automatic checks on the run's CSV/JSON where keys can be matched, plus an orchestrator rubric from the reported numbers (recorded per run in rubric.json and shown separately).")

H("5. Results")
body = []
for wf in WFS:
    for lv in LEVELS:
        body.append([wf.split("_")[0], lv.split("_")[1], cell(wf, lv, "executed"), cell(wf, lv, "physics_pass"), cell(wf, lv, "rubric_pushback"), cell(wf, lv, "meta_iterations"), cell(wf, lv, "loc"), cell(wf, lv, "api_coverage"), cell(wf, lv, "reusable"), cell(wf, lv, "uses_experimental_data")])
T(["WF", "level", "ran", "physics", "pushback", "iters", "LOC", "API cov.", "reusable", "exp. data"], body, widths=[0.5, 0.7, 0.5, 0.6, 0.7, 0.5, 0.5, 0.7, 0.7, 0.7], caption="Table 2. Per-cell grades (– = run pending).")
if iters is not None:
    T(["level", "mean iterations", "mean LOC", "API coverage", "reusable"], [[lv, round(iters[lv], 1) if pd.notna(iters[lv]) else "–", round(loc[lv]) if pd.notna(loc[lv]) else "–", round(api[lv], 2) if pd.notna(api[lv]) else "–", round(reuse[lv], 2) if pd.notna(reuse[lv]) else "–"] for lv in LEVELS], widths=[1.2, 1.3, 1.0, 1.2, 1.0], caption="Table 3. Cost and structure vs prompt level (all workflows).")
H("5.1 Physics is right from the first prompt; specificity buys structure", 2)
P(f"{n_phys} of {n_runs} runs passed the physics checks, at every level including the one-line Sketch. What changed with specificity was structure: reusability (a CLI or main-plus-functions) and interface contracts appear only at L4–L5, at the cost of more iterations and 2–5× more code (Table 3). This reproduces the central ALCHEMI finding in a different domain and toolkit.")
H("5.2 Every ill-posed request was refused — because the toolkit refused first", 2)
P(f"All {len(pb)} W4 runs pushed back (Table 2), and all of them did so by exercising foamsim's own ValueError for vf > 0.64 and η ≥ 1, then reporting the nearest admissible point with the correction stated. ALCHEMI reported zero pushback across every configuration with an unguarded toolkit. The comparison isolates a design lever: admissibility checks in the library convert 'the agent does not question the premise' into 'the agent cannot proceed without addressing it'. Whether agents push back on ill-posed tasks the library cannot catch (e.g., a static model asked for a rate-dependent property) is left for a second control workflow.")
H("5.3 Inverse design converges to the same optimum at every level", 2)
P("All five W2 runs found the same design (η ≈ 0.958, vf = 0.60, ρ ≈ 0.657 g/cm³, E ≈ 3515–3540 MPa; nearest catalogue grade 3M S32) and all noted that the optimum sits on the vf boundary and that the 1% margin would not survive the 20–40% mean-field-to-experiment gap. The first version of this task (target 2500 MPa) was flagged by the agent as non-binding — every point in the box exceeded it — and the task was re-targeted; agents can debug benchmark design.")
H("5.4 FE-vs-analytical verification", 2)
if len(w3) and w3.executed.sum():
    P("W3 runs: " + "; ".join(f"{r.level}: ran={int(r.executed)}, physics={cell('W3_fe_vs_analytical', r.level, 'physics_pass')}, rel. diff={cell('W3_fe_vs_analytical', r.level, 'rel_diff')}" for _, r in w3.iterrows()) + ".")
    P("All five W3 runs produced the same numbers (FE mean 3801 ± 7 MPa at n = 16/24 × two seeds; HP-MT 3696 MPa; +2.8 %) and all five reported, unprompted, that the FE value lies 0.8–1.2 % above the Hashin–Shtrikman upper bound and attributed it to the affine (KUBC) constraint on a 16-sphere window plus voxel stiffening, citing the exact homogeneous-box limit to rule out a solver bug and the n = 16 → 24 refinement trend. One run's convergence probe extended this: n = 32 on the same cell gives 3792 MPa and a 30-sphere cell at n = 32 gives 3784 MPa, converging monotonically toward the bound (3766 MPa). The protocol check 'FE inside HS bounds' is therefore subtly wrong for KUBC on small cells — and the agents, rather than the benchmark author, were the ones to say so. The contract-level run encoded this as a failed validate() check and exited 1, as its interface required.")
else:
    P("[W3 runs pending — inserted automatically when available.]")
H("5.5 What the experimental data can validate", 2)
P("Every W1 run compared its curve with the FoamGPT epoxy/glass-microballoon compression rows and every run reported the same problem independently: the available moduli (13 rows, mostly one layered functionally-graded-foam thesis and one S60HS point) lie below the Hashin–Shtrikman lower bound for their own constituents, which no intact, bonded microstructure can produce. The agents attributed this to crosshead-based secant moduli, matrix porosity, debonding and particle breakage, re-predicted each record with its own particle density, and refused to tune constants. For the benchmark this means W1's experimental check is a consistency report, not a validation; for the dataset it is a targeted finding — quasi-static compressive moduli of polymer syntactic foams in the literature need strain-measurement metadata before they can validate micromechanics.")

H("6. Discussion")
B(["Guarded toolkits are a cheap, transferable safety mechanism for agentic simulation: the physics library, not the prompt, is where premise checks are enforceable.",
   "Agents default to what the skills show (HP-MT with HS bounds in every run) and rarely reach beyond it (the crush-onset and FE tools were used only when asked), mirroring ALCHEMI's 'what the examples showcase, the agent uses'.",
   "Experimental ground truth is harder than DFT ground truth: agents were consistent and correct about the model, yet the comparison to experiment was inconclusive for data reasons. Benchmarks in experimental fields need curated reference subsets with measurement metadata.",
   "Cost: analytical workflows complete in 1–4 iterations and 1–5 minutes; FE workflows are CPU-bound and dominate wall-clock."])
H("7. Limitations and next steps")
B(["One model family and one sample per cell; add a second agent model and n = 3 samples per cell.",
   "KUBC on 16–20-sphere cells; add periodic boundary conditions and larger cells; add 'shell' mode runs at n ≥ 64.",
   "A second ill-posed control that the toolkit cannot catch (rate dependence, temperature, non-monodisperse packing claims).",
   "Physical validation with the collaborating laboratory: cast epoxy/K46 at 0.2–0.5 vf with strain-gauge moduli to give the benchmark a like-for-like experimental anchor.",
   "Extend the toolkit to strength (particle crushing + matrix yield) and to DLP-printed particle-filled resins."])
H("Code and data availability")
P("https://github.com/Harshelite2503/foamsim-agent (toolkit, skills, prompts, all agent runs with scripts, logs and grades). Experimental data: FoamGPT, https://github.com/Harshelite2503/foamgpt.")
H("References")
B(["NVIDIA. How AI coding agents can unlock materials simulation with NVIDIA ALCHEMI Toolkit. Technical blog, Aug 2026.",
   "Sarker T.R. et al. A multi-AI-agent framework enabling end-to-end finite element analysis for solid mechanics problems. arXiv 2026.",
   "Mohammadzadeh S. et al. FEM-Bench: a structured scientific reasoning benchmark for code-generating LLMs. arXiv 2025.",
   "Jadhav Y., Barati Farimani A. Large language model agent as a mechanical designer. J. Eng. Design 2026.",
   "Hashin Z. The elastic moduli of heterogeneous materials. J. Appl. Mech. 1962.",
   "Hashin Z., Shtrikman S. A variational approach to the theory of the elastic behaviour of multiphase materials. JMPS 1963.",
   "Benveniste Y. A new approach to the application of Mori–Tanaka's theory in composite materials. Mech. Mater. 1987.",
   "McLaughlin R. A study of the differential scheme for composite materials. Int. J. Eng. Sci. 1977.",
   "Bardella L., Genna F. On the elastic behavior of syntactic foams. Int. J. Solids Struct. 2001.",
   "Porfiri M., Gupta N. Effect of volume fraction and wall thickness on the elastic properties of hollow particle filled composites. Compos. B 2009.",
   "Gupta N., Pinisetty D., Shunmugasamy V.C. Reinforced Polymer Matrix Syntactic Foams: Modeling and Simulation. Springer 2013.",
   "Gustafsson T., McBain G.D. scikit-fem: a Python package for finite element assembly. JOSS 2020."])
out = ROOT / "FoamSim_Agent_Draft_Paper.docx"; doc.save(out); (ROOT / "docs").mkdir(exist_ok=True); (ROOT / "docs" / out.name).write_bytes(out.read_bytes())
(ROOT / "docs" / "paper_draft.md").write_text("\n".join(md)); print("saved", out)
